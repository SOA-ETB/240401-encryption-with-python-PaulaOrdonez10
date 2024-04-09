import os
from docx import Document
from docx.shared import Pt

try:
    # Get the current directory path
    current_directory = os.path.dirname(os.path.realpath(__file__))

    # Create Word document
    doc = Document()

    # Add the CV title with specific font and size
    title = doc.add_heading('Curriculum Vitae', level=1)
    run = title.runs[0]
    font = run.font
    font.name = 'Arial'
    font.size = Pt(18)

    # Personal Information
    doc.add_heading('Personal Information', level=2)

    # Name
    name = "Paula A. Ordoñez Montoya"
    doc.add_paragraph(f'Name: {name}')

    # Address
    address = "57th Granary Court, North Ring Road"
    doc.add_paragraph(f'Address: {address}')

    # Phone
    phone = "+353 830906817"
    doc.add_paragraph(f'Phone: {phone}')

    # Email
    email = "paulamadrid93@gmail.com"
    doc.add_paragraph(f'Email: {email}')

    # LinkedIn
    linkedin = "www.linkedin.com/in/paula-ordóñez-montoya"
    doc.add_paragraph(f'LinkedIn: {linkedin}')

    # Add a blank line
    doc.add_paragraph()

    # Work Experience
    doc.add_heading('Work Experience', level=2)

    # Define work experience
    experience = [
        {
            "role": "Trainee Software Process Automation Developer (RPA)",
            "company": "School of Automation",
            "date": "September 2023 - Present",
            "description": "1 Year Traineeship Programme with SOA and LCETB.\n"
                           "- RPA (Robotic Process Automation): RPA Lifecycle development, deployment, testing\n"
                           "- RPA Business Analysis: process capture, analysis, optimization and documentation\n"
                           "- UiPath Studio, Orchestrator, AI Center\n"
                           "- Blue Prism\n"
                           "- Microsoft Power Platform: Power Automate, Power Apps\n"
                           "- SQL: Designing, Building and Query databases with MS Access\n"
                           "- Python: Core software development concepts\n"
                           "- ELK / Kibana: Data Analysis and Visualization"
        },
        {
            "role": "Shift Supervisor",
            "company": "Starbucks",
            "date": "September 2022 - Present",
            "description": "- Partner of the quarter\n"
                           "- Supervisor of shift\n"
                           "- Customer service\n"
                           "- Barista training\n"
                           "- Keep the floor area clean\n"
                           "- Coffee knowledge\n"
                           "- Kitchen porter knowledge\n"
                           "- Money handling"
        },
        {
            "role": "General Manager Customer Service",
            "company": "Mi Store Spain",
            "date": "February 2017 - May 2019",
            "description": "- Ability to solve problems as efficiently and quickly as possible.\n"
                           "- Control of stores for claims\n"
                           "- Store inventories: electronically and physically\n"
                           "- Carry out protocols, from the management of the shops to the technical service\n"
                           "- Using initiative to identify tasks that need to be completed\n"
                           "- Having the drive to work hard and contribute to the success of each store\n"
                           "- A friendly communication style with every team and customers\n"
                           "- Make budgets and invoices using SAGE, Polarik (Company software) and Excel 2019."
        },
        {
            "role": "Sales Assistant",
            "company": "HellermannTyton Global",
            "date": "July 2016 - October 2016",
            "description": "- Make a product budget\n"
                           "- Offer products through different channels including mail and telephone\n"
                           "- Promotions through newsletter"
        }
    ]

    # Add work experience details
    for job in experience:
        p = doc.add_paragraph()
        p.add_run(f'{job["role"]} - {job["company"]}, {job["date"]}').bold = True
        p = doc.add_paragraph(job["description"])

    # Add a blank line
    doc.add_paragraph()

    # Education
    doc.add_heading('Education', level=2)

    # Define education
    education = [
        {
            "title": "Robotic Process Automation Developer",
            "institution": "Limerick and Clare Education and Training Board",
            "date": "June 2023 - September 2023"
        },
        {
            "title": "ESOL, English Level 1",
            "institution": "Birmingham Metropolitan College",
            "date": "2019 - 2020"
        },
        {
            "title": "General English Course, Certificate B1",
            "institution": "Cork English World",
            "date": "2019 - 2019"
        },
        {
            "title": "Bachelor's Degree in Business Administration",
            "institution": "Complutense University of Madrid",
            "date": "2013 - 2016"
        }
       
    ]

    # Add education details
    for item in education:
        p = doc.add_paragraph()
        p.add_run(f'{item["title"]} - {item["institution"]}, {item["date"]}').bold = True

    # Add a blank line
    doc.add_paragraph()

    # Skills
    doc.add_heading('Skills', level=2)

    # Define skills
    skills = ["Blue Prism", "UiPath", "HTML", "Python"]

    # Add skills
    for skill in skills:
        doc.add_paragraph(skill)

    # Save document as "Paula_cv.docx" in the current directory
    file_path = os.path.join(current_directory, "Paula_cv.docx")
    doc.save(file_path)

    print(f'Curriculum Vitae created successfully as "{file_path}"')

except Exception as e:
    print(f'An error occurred: {e}')
