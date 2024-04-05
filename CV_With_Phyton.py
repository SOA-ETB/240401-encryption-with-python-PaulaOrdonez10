from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()

# Add the CV title with specific font and size
title = doc.add_heading('Curriculum Vitae', level=1)
run = title.runs[0]
font = run.font
font.name = 'Arial'
font.size = Pt(18)

# Personal Information
doc.add_heading('Personal Information', level=2)

# Name
name = "Paula A. Ordoñez Montoya"
doc.add_paragraph(f'Name: {name}')

# Address
address = "57th Granary Court, North Ring Road"
doc.add_paragraph(f'Address: {address}')

# Phone
phone = "+353 830906817"
doc.add_paragraph(f'Phone: {phone}')

# Email
email = "paulamadrid93@gmail.com"
doc.add_paragraph(f'Email: {email}')

# LinkedIn
linkedin = "www.linkedin.com/in/paula-ordóñez-montoya"
doc.add_paragraph(f'LinkedIn: {linkedin}')

# Add a blank line
doc.add_paragraph()

# Work Experience
doc.add_heading('Work Experience', level=2)

# Define work experience
experiencia = [
    {
        "role": "Trainee Software Process Automation Developer (RPA)",
        "company": "SOA",
        "date": "September 2023 - Present",
        "description": "1 Year Traineeship Programme with SOA and LCETB."
    },
    {
        "role": "Shift Supervisor",
        "company": "Starbucks",
        "date": "September 2022 - March 2023",
        "description": "Partner of the quarter."
    },
    # Add the rest of the work experience here
]

# Add work experience details
for job in experiencia:
    p = doc.add_paragraph()
    p.add_run(f'{job["role"]} - {job["company"]}, {job["date"]}').bold = True
    p = doc.add_paragraph(job["description"])

# Add a blank line
doc.add_paragraph()

# Education
doc.add_heading('Education', level=2)

# Define education
educacion = [
    {
        "title": "Robotic Process Automation Developer",
        "institution": "Limerick and Clare Education and Training Board",
        "date": "June 2023 - September 2023"
    },
    {
        "title": "ESOL, English Level 1",
        "institution": "Birmingham Metropolitan College",
        "date": "2019 - 2020"
    },
    # Add the rest of the education here
]

# Add education details
for education in educacion:
    p = doc.add_paragraph()
    p.add_run(f'{education["title"]} - {education["institution"]}, {education["date"]}').bold = True

# Add a blank line
doc.add_paragraph()

# Skills
doc.add_heading('Skills', level=2)

# Define skills
skills = ["Blue Prism", "UiPath", "HTML"]

# Skills
for skill in skills:
    doc.add_paragraph(skill)

# Save document as "cv.docx"
doc.save('cv.docx')

print('Curriculum Vitae created successfully as "cv.docx"')
